#!/usr/bin/env python
# coding: utf-8
"""
Generate a Word CV (files/cv.docx) from structured site data.

Single source of truth:
  * _data/cv.yml         -> education, experience, skills, header
  * _publications/*.md   -> Publications  (uses the AMA `citation` field)
  * _talks/*.md          -> Presentations
  * _teaching/*.md       -> Teaching

Run locally:
    pip install python-docx pyyaml
    python scripts/build_cv.py

The GitHub Actions workflow runs this, then converts the .docx to
files/cv.pdf with LibreOffice, before building the Jekyll site. Both files
are published, so the "Download PDF / Word" buttons on /cv/ always serve a
document regenerated from the same data that renders the web page.

Requires: python-docx, pyyaml
"""

import os
import re
import glob
import datetime

import yaml
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(HERE)
DATA = os.path.join(REPO, "_data", "cv.yml")
OUT = os.path.join(REPO, "files", "cv.docx")

ACCENT = RGBColor(0x1A, 0x5B, 0x8A)   # heading color
DARK = RGBColor(0x22, 0x22, 0x22)


# --------------------------------------------------------------------------
# front-matter helpers
# --------------------------------------------------------------------------
def read_front_matter(path):
    txt = open(path, encoding="utf-8").read()
    m = re.match(r"^---\n(.*?)\n---", txt, re.S)
    return yaml.safe_load(m.group(1)) if m else {}


def collection(name):
    items = []
    for fn in glob.glob(os.path.join(REPO, name, "*.md")):
        fm = read_front_matter(fn)
        if fm:
            items.append(fm)
    return items


def as_date(v):
    if isinstance(v, datetime.date):
        return v
    if isinstance(v, str):
        for fmt in ("%Y-%m-%d", "%Y-%m", "%Y"):
            try:
                return datetime.datetime.strptime(v[:len(fmt) + 2], fmt).date()
            except ValueError:
                continue
    return datetime.date(1900, 1, 1)


# --------------------------------------------------------------------------
# docx styling helpers
# --------------------------------------------------------------------------
def add_runs_with_bold(paragraph, text, bold_pattern=r"<strong>(.*?)</strong>",
                       also_bold=None, size=10):
    """Add `text` to paragraph, rendering <strong>..</strong> (and optional
    `also_bold` substrings, e.g. 'Benson') in bold."""
    # normalize markdown ** ** to the strong tag, then split on strong.
    text = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", text)
    parts = re.split(bold_pattern, text)
    # re.split with one capturing group: even idx = normal, odd idx = bold
    for i, chunk in enumerate(parts):
        if not chunk:
            continue
        if i % 2 == 1:
            r = paragraph.add_run(chunk)
            r.bold = True
            r.font.size = Pt(size)
        else:
            if also_bold:
                _add_with_keyword_bold(paragraph, chunk, also_bold, size)
            else:
                r = paragraph.add_run(chunk)
                r.font.size = Pt(size)


def _add_with_keyword_bold(paragraph, text, keyword, size):
    """Bold any run of text containing `keyword` surname (e.g. 'Benson FM')."""
    # bold the surname + following initials token, e.g. "Benson JS"
    pat = re.compile(r"(" + re.escape(keyword) + r"[ ,]?[A-Z]{0,3}\.?)")
    for i, chunk in enumerate(pat.split(text)):
        if not chunk:
            continue
        r = paragraph.add_run(chunk)
        r.font.size = Pt(size)
        if i % 2 == 1:
            r.bold = True


def section_heading(doc, title):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A5B8A")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def bullet(doc, text, level=0, size=10, italic_label=False):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    add_runs_with_bold(p, text, size=size)
    return p


def plain(doc, text, size=10, bold=False, italic=False, space_after=2,
          align=None, also_bold=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    add_runs_with_bold(p, text, size=size, also_bold=also_bold)
    for r in p.runs:
        if bold:
            r.bold = True
        if italic:
            r.italic = True
    return p


# --------------------------------------------------------------------------
# document sections
# --------------------------------------------------------------------------
def build_header(doc, cv):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(cv["name"])
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(0)

    if cv.get("headline"):
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = h.add_run(cv["headline"])
        hr.font.size = Pt(11)
        hr.italic = True
        h.paragraph_format.space_after = Pt(2)

    c = cv.get("contact", {})
    bits = []
    if c.get("email"):
        bits.append(c["email"])
    if c.get("website"):
        bits.append(c["website"].replace("https://", "").replace("http://", ""))
    if c.get("orcid"):
        bits.append("ORCID: " + c["orcid"])
    if bits:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run("  •  ".join(bits))
        cr.font.size = Pt(9.5)
        cp.paragraph_format.space_after = Pt(4)


def build_education(doc, cv):
    section_heading(doc, "Education")
    for ed in cv.get("education", []):
        head = ed["institution"] + (f", {ed['location']}" if ed.get("location") else "")
        plain(doc, head, bold=True, size=10.5, space_after=0)
        for deg in ed.get("degrees", []):
            plain(doc, deg, size=10, space_after=0)
        if ed.get("note"):
            plain(doc, ed["note"], italic=True, size=9.5, space_after=0)
        for d in ed.get("details", []):
            bullet(doc, f"<strong>{d['label']}:</strong> {d['text']}", size=9.5)


def build_experience(doc, cv):
    section_heading(doc, "Experience")
    for job in cv.get("experience", []):
        line = f"{job['org']}: {job['role']}"
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        lr = p.add_run(line)
        lr.bold = True
        lr.font.size = Pt(10.5)
        if job.get("location") or job.get("dates"):
            meta = " — ".join([x for x in [job.get("location"), job.get("dates")] if x])
            mp = doc.add_paragraph()
            mp.paragraph_format.space_after = Pt(0)
            mr = mp.add_run(meta + (f"  ({job['note']})" if job.get("note") else ""))
            mr.italic = True
            mr.font.size = Pt(9.5)
        for b in job.get("bullets", []):
            bullet(doc, b, size=10)


def build_skills(doc, cv):
    section_heading(doc, "Skills & Interests")
    for s in cv.get("skills", []):
        bullet(doc, f"<strong>{s['label']}:</strong> {s['text']}", size=10)


def build_publications(doc):
    pubs = sorted(collection("_publications"), key=lambda x: as_date(x.get("date")),
                  reverse=True)
    if not pubs:
        return
    section_heading(doc, "Publications")
    for i, p in enumerate(pubs, 1):
        cite = p.get("citation") or p.get("title", "")
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.first_line_indent = Inches(-0.25)
        num = para.add_run(f"{i}. ")
        num.font.size = Pt(9.5)
        add_runs_with_bold(para, cite, also_bold="Benson", size=9.5)


def build_presentations(doc):
    talks = sorted(collection("_talks"), key=lambda x: as_date(x.get("date")),
                   reverse=True)
    if not talks:
        return
    section_heading(doc, "Presented Abstracts & Presentations")
    for i, t in enumerate(talks, 1):
        title = t.get("title", "").strip()
        meta_bits = []
        if t.get("type"):
            meta_bits.append(t["type"])
        venue_loc = ", ".join([x for x in [t.get("venue"), t.get("location")] if x])
        d = as_date(t.get("date"))
        date_str = d.strftime("%B %Y") if d.year > 1900 else ""
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.first_line_indent = Inches(-0.25)
        num = para.add_run(f"{i}. ")
        num.font.size = Pt(9.5)
        tr = para.add_run(title + ". ")
        tr.bold = True
        tr.font.size = Pt(9.5)
        tail = ". ".join([b for b in [", ".join(meta_bits + [venue_loc]).strip(", ")] if b])
        if date_str:
            tail = (tail + ", " + date_str) if tail else date_str
        if t.get("authors"):
            add_runs_with_bold(para, str(t["authors"]) + ". ", also_bold="Benson", size=9.5)
        if tail:
            er = para.add_run(tail + ".")
            er.font.size = Pt(9.5)


def build_teaching(doc):
    teach = sorted(collection("_teaching"), key=lambda x: as_date(x.get("date")),
                   reverse=True)
    if not teach:
        return
    section_heading(doc, "Teaching")
    for t in teach:
        venue_loc = ", ".join([x for x in [t.get("venue"), t.get("location")] if x])
        d = as_date(t.get("date"))
        date_str = d.strftime("%B %Y") if d.year > 1900 else ""
        line = t.get("title", "")
        meta = " — ".join([x for x in [t.get("type"), venue_loc, date_str] if x])
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        tr = p.add_run(line + ". ")
        tr.bold = True
        tr.font.size = Pt(9.5)
        if meta:
            mr = p.add_run(meta + ".")
            mr.font.size = Pt(9.5)


# --------------------------------------------------------------------------
def main():
    cv = yaml.safe_load(open(DATA, encoding="utf-8"))
    doc = Document()

    # base style + tight margins
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    for sec in doc.sections:
        sec.top_margin = Inches(0.5)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin = Inches(0.6)
        sec.right_margin = Inches(0.6)

    build_header(doc, cv)
    build_education(doc, cv)
    build_experience(doc, cv)
    build_skills(doc, cv)
    build_publications(doc)
    build_presentations(doc)
    build_teaching(doc)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
